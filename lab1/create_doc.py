from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Function to add code block style
def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    # Make it look like a code box
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), "D9EAF7")  # light blue background
    p._element.get_or_add_pPr().append(shading_elm)

# Create a new Document
doc = Document()

# Title
doc.add_heading('House Price Dataset Analysis', level=1)

# ---------------- Sample Data ----------------
doc.add_heading('1. Sample Data (First 5 Rows)', level=2)
add_code_block(doc, "print(df.head())")

sample_data = [
    ["price", "area", "bedrooms", "bathrooms", "stories", "mainroad", "guestroom", "basement",
     "hotwaterheating", "airconditioning", "parking", "prefarea", "furnishingstatus"],
    [13300000, 7420, 4, 2, 3, "yes", "no", "no", "no", "yes", 2, "yes", "furnished"],
    [12250000, 8960, 4, 4, 4, "yes", "no", "no", "no", "yes", 3, "no", "furnished"],
    [12250000, 9960, 3, 2, 2, "yes", "no", "yes", "no", "no", 2, "yes", "semi-furnished"],
    [12215000, 7500, 4, 2, 2, "yes", "no", "yes", "no", "yes", 3, "yes", "furnished"],
    [11410000, 7420, 4, 1, 2, "yes", "yes", "yes", "no", "yes", 2, "no", "furnished"],
]

table = doc.add_table(rows=len(sample_data), cols=len(sample_data[0]))
table.style = 'Light List Accent 1'
for i, row in enumerate(sample_data):
    for j, val in enumerate(row):
        table.cell(i, j).text = str(val)

# ---------------- Range ----------------
doc.add_heading('2. Range (Max – Min)', level=2)
add_code_block(doc, "range_values = df.max(numeric_only=True) - df.min(numeric_only=True)\nprint('Range:\\n', range_values, '\\n')")

range_data = [
    ["Feature", "Range"],
    ["price", 11550000],
    ["area", 14550],
    ["bedrooms", 5],
    ["bathrooms", 3],
    ["stories", 3],
    ["parking", 3]
]
table = doc.add_table(rows=len(range_data), cols=len(range_data[0]))
table.style = 'Light Grid Accent 2'
for i, row in enumerate(range_data):
    for j, val in enumerate(row):
        table.cell(i, j).text = str(val)

# ---------------- Q1 ----------------
doc.add_heading('3. Q1 (25th Percentile)', level=2)
add_code_block(doc, "q1 = df.quantile(0.25, numeric_only=True)\nprint('Q1 (25th percentile):\\n', q1, '\\n')")

q1_data = [
    ["Feature", "Q1 Value"],
    ["price", 3430000],
    ["area", 3600],
    ["bedrooms", 2],
    ["bathrooms", 1],
    ["stories", 1],
    ["parking", 0]
]
table = doc.add_table(rows=len(q1_data), cols=len(q1_data[0]))
table.style = 'Light Grid Accent 2'
for i, row in enumerate(q1_data):
    for j, val in enumerate(row):
        table.cell(i, j).text = str(val)

# ---------------- Q3 ----------------
doc.add_heading('4. Q3 (75th Percentile)', level=2)
add_code_block(doc, "q3 = df.quantile(0.75, numeric_only=True)\nprint('Q3 (75th percentile):\\n', q3, '\\n')")

q3_data = [
    ["Feature", "Q3 Value"],
    ["price", 5740000],
    ["area", 6360],
    ["bedrooms", 3],
    ["bathrooms", 2],
    ["stories", 2],
    ["parking", 1]
]
table = doc.add_table(rows=len(q3_data), cols=len(q3_data[0]))
table.style = 'Light Grid Accent 2'
for i, row in enumerate(q3_data):
    for j, val in enumerate(row):
        table.cell(i, j).text = str(val)

# ---------------- IQR ----------------
doc.add_heading('5. Interquartile Range (IQR)', level=2)
add_code_block(doc, "iqr = q3 - q1\nprint('Interquartile Range (IQR):\\n', iqr, '\\n')")

iqr_data = [
    ["Feature", "IQR"],
    ["price", 2310000],
    ["area", 2760],
    ["bedrooms", 1],
    ["bathrooms", 1],
    ["stories", 1],
    ["parking", 1]
]
table = doc.add_table(rows=len(iqr_data), cols=len(iqr_data[0]))
table.style = 'Light Grid Accent 2'
for i, row in enumerate(iqr_data):
    for j, val in enumerate(row):
        table.cell(i, j).text = str(val)

import pandas as pd
import seaborn as sns
import matplotlib.pyplot as plt
from scipy.stats import chi2_contingency, pearsonr, probplot
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

# ---------- Load dataset ----------
df = pd.read_csv("Housing.csv")   # Make sure Housing.csv is in the same folder

# ---------- Helper to format code blocks ----------
def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), "D9EAF7")  # light blue background
    p._element.get_or_add_pPr().append(shading_elm)

# ---------- Create Word Document ----------
doc = Document()
doc.add_heading("Statistical & Graphical Analysis of Housing Dataset", level=1)

# ---------- Boxplots ----------
doc.add_heading("1. Box Plots", level=2)
code_boxplot = """plt.figure(figsize=(10, 6))
sns.boxplot(data=df['price'])
plt.title('Box Plot for Price')

sns.boxplot(data=df.drop(columns=['price','area']))
plt.title('Box Plot (excluding price and area'))
plt.show()"""
add_code_block(doc, code_boxplot)

plt.figure(figsize=(10, 6))
sns.boxplot(data=df['price'])
plt.title("Box Plot for Price")
plt.savefig("boxplot_price.png")
plt.close()

plt.figure(figsize=(10, 6))
sns.boxplot(data=df.drop(columns=['price', 'area']))
plt.title("Box Plot (excluding price and area)")
plt.savefig("boxplot_features.png")
plt.close()

doc.add_picture("boxplot_price.png", width=Inches(5))
doc.add_picture("boxplot_features.png", width=Inches(5))

# ---------- QQ Plot ----------
doc.add_heading("2. QQ Plot for Price", level=2)
code_qq = """stats.probplot(df["price"].dropna(), dist="norm", plot=plt)
plt.title("QQ Plot — price")"""
add_code_block(doc, code_qq)

plt.figure(figsize=(6, 6))
probplot(df["price"].dropna(), dist="norm", plot=plt)
plt.title("QQ Plot — price")
plt.tight_layout()
plt.savefig("qqplot_price.png")
plt.close()
doc.add_picture("qqplot_price.png", width=Inches(4.5))

# ---------- Histogram ----------
doc.add_heading("3. Histogram for Price", level=2)
code_hist = """sns.histplot(df["price"], bins=30, kde=True)
plt.title("Histogram — price")"""
add_code_block(doc, code_hist)

plt.figure(figsize=(8, 4))
sns.histplot(df["price"], bins=30, kde=True)
plt.title("Histogram — price")
plt.xlabel("price")
plt.ylabel("count")
plt.tight_layout()
plt.savefig("hist_price.png")
plt.close()
doc.add_picture("hist_price.png", width=Inches(5))

# ---------- Chi-square Test ----------
doc.add_heading("4. Chi-Square Test for Price Bin vs Furnishing Status", level=2)
code_chi2 = """ct = pd.crosstab(df["price_bin"], df["furnishingstatus"])
chi2, p, dof, expected = chi2_contingency(ct.values)
print(ct)
print("Chi-square:", chi2, "| p-value:", p, "| dof:", dof)
print("Expected frequencies:\\n", expected)"""
add_code_block(doc, code_chi2)

df["furnishingstatus"] = df["furnishingstatus"].astype("category")
df["price_bin"] = pd.qcut(df["price"], q=3, labels=["low","medium","high"], duplicates="drop")
ct = pd.crosstab(df["price_bin"], df["furnishingstatus"])
chi2, p, dof, expected = chi2_contingency(ct.values)

doc.add_paragraph("Contingency Table:")
table = doc.add_table(rows=1+ct.shape[0], cols=1+ct.shape[1])
table.style = "Light Grid Accent 2"
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "price_bin"
for j, col in enumerate(ct.columns):
    hdr_cells[j+1].text = str(col)
for i, idx in enumerate(ct.index):
    row = table.rows[i+1].cells
    row[0].text = str(idx)
    for j, col in enumerate(ct.columns):
        row[j+1].text = str(ct.loc[idx, col])

doc.add_paragraph(f"\nChi-square = {chi2:.4f}, p-value = {p:.4e}, dof = {dof}")
doc.add_paragraph("Expected frequencies:")
table2 = doc.add_table(rows=expected.shape[0], cols=expected.shape[1])
table2.style = "Light Grid Accent 2"
for i in range(expected.shape[0]):
    for j in range(expected.shape[1]):
        table2.cell(i,j).text = f"{expected[i,j]:.2f}"
doc.add_paragraph("Decision: " + ("Reject H0 (dependent)" if p <= 0.05 else "Fail to reject H0 (independent)"))

# ---------- Pearson Correlation ----------
doc.add_heading("5. Pearson Correlation", level=2)
code_pearson = """r_pa, p_pa = pearsonr(df["price"], df["area"])
print("Pearson r (price, area) =", r_pa, ", p-value =", p_pa)

corr_mat = df[["price","area","bedrooms","bathrooms","stories","parking"]].corr(method="pearson")
print(corr_mat)"""
add_code_block(doc, code_pearson)

r_pa, p_pa = pearsonr(df["price"].astype(float), df["area"].astype(float))
doc.add_paragraph(f"Pearson r (price, area) = {r_pa:.4f}, p-value = {p_pa:.4f}")

corr_mat = df[["price","area","bedrooms","bathrooms","stories","parking"]].corr(method="pearson")
doc.add_paragraph("Correlation Matrix:")
table3 = doc.add_table(rows=1+corr_mat.shape[0], cols=1+corr_mat.shape[1])
table3.style = "Light Grid Accent 2"
hdr = table3.rows[0].cells
hdr[0].text = ""
for j, col in enumerate(corr_mat.columns):
    hdr[j+1].text = col
for i, idx in enumerate(corr_mat.index):
    row = table3.rows[i+1].cells
    row[0].text = str(idx)
    for j, col in enumerate(corr_mat.columns):
        row[j+1].text = f"{corr_mat.loc[idx, col]:.4f}"

# ---------- Save Word File ----------
doc.save("analysis_report.docx")
print("✅ Word document cr")


# ---------- Save Word File ----------
doc.save("analysis_report.docx")
print("✅ Word document created: analysis_report.docx")



